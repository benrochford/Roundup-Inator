import streamlit as st
import requests
import pandas as pd
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
import docx
import io
import markdown
import time
from langchain_community.llms import Ollama
from datetime import datetime, timedelta

DEFAULT_ROUNDUP_PROMPT = """Create a research round-up following this structure:

# ((apparent topic)) Research Round-up: ((apparent date range from papers))

Overview Paragraph with:
- 3-4 major themes identified from the papers
- Brief (1-2 sentence) explanation of why each theme matters
- (Include inline links to papers that are most relevant to each theme)

## Notable Papers
IMPORTANT: You MUST discuss AT LEAST 5 papers in this section, preferably 6-8. Do not stop at 3 papers.
For each paper:
**Paper Title** (with link to Semantic Scholar URL)
Key findings in 2-3 sentences, 1 sentence of why it matters

## Quick Takes
- 8-15 highlights from other interesting papers
- Focus on actionable insights or surprising findings
- Format like "a study by AUTHOR et al. found that..." which is hyperlinked to the semanticscholar url of the paper

## Emerging Trends, Future Directions
in paragraph form:
- 5 or so important emerging trends based on developments from the research                                                     
- 3 sentence discussion of potential implications, future research directions that are coming"""

LOADING_MESSAGES = [
    "Reading abstracts",
    "Highlighting important bits",
    "Taking a snack break",
    "Checking Mastodon for hot takes",
    "Drawing connections",
    "Looking at pictures of national parks",
    "Procrastinating productively",
    "Making coffee",
    "Organizing sticky notes",
    "Double-checking citations",
    "Contemplating the nature of knowledge",
    "Doing desk stretches",
    "Looking at mildly interesting things on Reddit",
    "Petting nearby cats for inspiration",
    "Reorganizing browser bookmarks",
    "Watching one more YouTube video",
]

st.set_page_config(page_title="Research Roundup Generator", layout="wide")

# Add custom CSS for dark mode styling
st.markdown("""
    <style>
    .stApp {
        background-color: #1a1a1a;
        color: #ffffff;
    }
    .main {
        max-width: 1200px;
        padding: 2rem;
    }
    .stButton button {
        background-color: #c71585;
        color: white !important;
        border-radius: 4px;
        border: none;
        padding: 0.5rem 1rem;
    }
    .stButton button:hover {
        background-color: #ff4500;
        color: white !important;
    }
    .stButton button:active {
        color: white !important;
    }
    .stTextInput input {
        background-color: #2d2d2d;
        color: #ffffff;
        border-radius: 4px;
        border: 1px solid #404040;
    }
    .stMarkdown {
        color: #ffffff;
    }
    </style>
    """, unsafe_allow_html=True)

# Initialize session state
if 'papers_df' not in st.session_state:
    st.session_state.papers_df = None
if 'newsletter_content' not in st.session_state:
    st.session_state.newsletter_content = None
if 'custom_prompt' not in st.session_state:
    st.session_state.custom_prompt = DEFAULT_ROUNDUP_PROMPT
if 'openai_api_key' not in st.session_state:
    st.session_state.openai_api_key = None
if 'top_n_papers' not in st.session_state:
    st.session_state.top_n_papers = 35

def suggest_search_terms(topic_description, openai_api_key=None, use_ollama=False):
    if use_ollama:
        llm = Ollama(model="llama3.1")
    else:
        llm = ChatOpenAI(api_key=openai_api_key, model="gpt-4o")
    prompt = ChatPromptTemplate.from_template("""
    Given this research topic description, suggest 2-4 specific search terms that would be good for finding relevant papers on Semantic Scholar. Use Semantic Scholar's advanced search operators to create precise queries:

    Available operators:
    - Use quotes for exact phrases: "machine learning"
    - Use + to require terms: +ai +ethics
    - Use - to exclude terms: -medicine
    - Use | for OR: (ai | artificial intelligence)
    - Use * for prefix matching: neural*
    - Use ~N for fuzzy matching: algorithm~2
    - Use "word1 word2" ~N for proximity search

    Example good queries:
    - ((cloud computing) | virtualization) +security -medicine
    - "red blood cell" + artificial intelligence
    - "machine learning" ~3 healthcare
    - neuro* +cognition -psychology

    Topic description: {topic}
    
    Return only the search terms, one per line, nothing else. Make the queries specific and targeted using the operators above.
    """)
    
    messages = prompt.format_messages(topic=topic_description)
    response = llm.invoke(messages)
    
    # Handle both string (Ollama) and Message object (OpenAI) responses
    content = response if isinstance(response, str) else response.content
    
    # Clean up the response
    content = content.replace('```', '').strip()
    
    # Return cleaned terms
    return [term.strip() for term in content.split('\n') if term.strip()]

# SemanticScholar API setup
def collect_papers(queries, start_year=None, end_year=None, last_two_weeks=False, api_key=None, excluded_terms=None):
    url = "https://api.semanticscholar.org/graph/v1/paper/search/"
    params = {
        "fields": "title,url,abstract,citationCount,authors,publicationTypes,publicationDate,openAccessPdf",
        "limit": 100,
    }
    
    # Add date filtering
    if last_two_weeks:
        end_date = pd.Timestamp.now()
        start_date = end_date - pd.Timedelta(days=14)
        params["publicationDateOrYear"] = f"{start_date.strftime('%Y-%m-%d')}:{end_date.strftime('%Y-%m-%d')}"
    elif start_year and end_year:
        params["year"] = f"{start_year}-{end_year}"
    elif start_year:
        params["year"] = f"{start_year}-"
    elif end_year:
        params["year"] = f"-{end_year}"
    
    headers = {"x-api-key": api_key} if api_key else {}
    all_recent_papers = []
    results_by_term = {}  # Track results for each term
    
    for query in queries:
        # Don't wrap the query in quotes - let users specify their own syntax
        if excluded_terms:
            # Add excluded terms with proper syntax
            excluded_query = ' '.join(f'-{term}' for term in excluded_terms)
            query = f'{query} {excluded_query}'
        
        params["query"] = query
        with st.spinner(f"Searching for: {query}"):
            response = requests.get(url, params=params, headers=headers)
            papers = response.json().get('data', [])
            results_by_term[query] = len(papers)
            
            for paper in papers:
                paper['search_term'] = query
            
            all_recent_papers.extend(papers)
            # Show immediate results for this term
            st.write(f"📍 Found {len(papers)} papers for '{query}'")

        time.sleep(.5)  # avoid rate limiting
    
    return all_recent_papers, results_by_term

# Clean and process papers
def process_papers(papers):
    df = pd.DataFrame(papers)
    
    expected_columns = {
        'title': None,
        'url': None,
        'paperId': None,
        'abstract': None,
        'citationCount': 0,
        'authors': None,
        'publicationTypes': None,
        'publicationDate': None,
        'openAccessPdf': None,
        'search_term': None
    }
    
    # Add missing columns with default values
    for col, default in expected_columns.items():
        if col not in df.columns:
            df[col] = default
    
    # Process dates safely
    df['publicationDate'] = pd.to_datetime(df['publicationDate'], errors='coerce')
    df['year'] = df['publicationDate'].dt.year.fillna(-1).astype('Int64')
    df['month'] = df['publicationDate'].dt.month.fillna(-1).astype('Int64')
    
    # Process other columns safely
    df['openAccessUrl'] = df['openAccessPdf'].apply(lambda x: x['url'] if isinstance(x, dict) and 'url' in x else None)
    df['authors'] = df['authors'].apply(lambda x: [author['name'] for author in x] if isinstance(x, list) else None)
    
    # Reorder columns
    columns_order = ['title', 'url', 'paperId', 'abstract', 'citationCount', 'authors', 'publicationTypes', 
                     'publicationDate', 'year', 'month', 'openAccessUrl', 'search_term']
    
    # Only include columns that exist
    existing_columns = [col for col in columns_order if col in df.columns]
    return df[existing_columns].copy()

# Newsletter generation
def generate_newsletter(papers_df, openai_api_key=None, custom_prompt=None, use_ollama=False, top_n_papers=35):
    # Sort by both date and citation count (giving more weight to recent papers)
    papers_df['score'] = papers_df['citationCount'].fillna(0) + (1 / (1 + (pd.Timestamp.now() - papers_df['publicationDate']).dt.days))
    
    # Get top N papers based on score
    papers_df = papers_df.sort_values('score', ascending=False).head(top_n_papers)
    
    newsletter_prompt = ChatPromptTemplate.from_template(
        custom_prompt + "\n\nPapers to analyze:\n{topic_analyses}"
    )
    
    # Format papers for prompt
    topic_analyses = []
    for topic in papers_df['search_term'].unique():
        topic_papers = papers_df[papers_df['search_term'] == topic]
        
        papers_text = []
        for _, paper in topic_papers.iterrows():
            authors = ', '.join(author for author in paper['authors']) if isinstance(paper['authors'], list) else 'Unknown'
            paper_text = f"""Title: {paper['title']}
            URL: {paper['url']}
            Authors: {authors}
            Abstract: {paper['abstract']}
            Publication Date: {paper['publicationDate']}
            """
            papers_text.append(paper_text)
        
        topic_text = f"Topic: {topic}\n" + "\n".join(papers_text)
        topic_analyses.append(topic_text)
    
    topic_analyses_text = "\n---\n".join(topic_analyses)
    
    # Generate newsletter
    if use_ollama:
        llm = Ollama(model="llama3.1")
    else:
        llm = ChatOpenAI(api_key=openai_api_key, model="gpt-4o")
    messages = newsletter_prompt.format_messages(topic_analyses=topic_analyses_text)
    newsletter = llm.invoke(messages)
    
    # Handle both Ollama (string) and OpenAI (Message object) responses
    return newsletter if isinstance(newsletter, str) else newsletter.content


def convert_to_docx(markdown_content):
    doc = docx.Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = docx.shared.Pt(11)
    
    # Customize heading styles
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Calibri'
    h1_style.font.size = docx.shared.Pt(16)
    h1_style.font.bold = True
    
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Calibri'
    h2_style.font.size = docx.shared.Pt(14)
    h2_style.font.bold = True
    
    # Convert markdown to HTML first for better parsing
    html = markdown.markdown(markdown_content)
    
    # Split content by sections
    sections = html.split('<h')
    
    for section in sections:
        if not section.strip():
            continue
            
        # Reconstruct the h tag for processing
        if section[0] != '<':
            section = '<h' + section
            
        # Process headings
        if '<h1>' in section:
            text = section.split('</h1>')[0].replace('<h1>', '')
            doc.add_heading(text, level=1)
            content = section.split('</h1>')[1]
        elif '<h2>' in section:
            text = section.split('</h2>')[0].replace('<h2>', '')
            doc.add_heading(text, level=2)
            content = section.split('</h2>')[1]
        else:
            content = section
            
        # Process paragraphs and lists
        paragraphs = content.split('<p>')
        for p in paragraphs:
            if not p.strip():
                continue
                
            # Handle lists
            if '<ul>' in p:
                items = p.split('<li>')
                for item in items[1:]:  # Skip first empty item
                    item_text = item.split('</li>')[0]
                    # Handle links within list items
                    item_text = process_links(item_text)
                    doc.add_paragraph(item_text, style='List Bullet')
            else:
                # Handle regular paragraphs
                p_text = p.replace('</p>', '').strip()
                if p_text:
                    # Handle links within paragraphs
                    p_text = process_links(p_text)
                    doc.add_paragraph(p_text)
    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def process_links(text):
    # Handle both markdown and HTML links
    import re
    
    # First handle markdown links [text](url)
    markdown_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    
    def replace_markdown_link(match):
        text, url = match.groups()
        return f"{text} ({url})"
    
    # Then handle HTML links <a href="url">text</a>
    html_pattern = r'<a\s+href="([^"]+)"[^>]*>([^<]+)</a>'
    
    def replace_html_link(match):
        url, text = match.groups()
        return f"{text} ({url})"
    
    # Apply both replacements
    text = re.sub(markdown_pattern, replace_markdown_link, text)
    text = re.sub(html_pattern, replace_html_link, text)
    
    # Clean up any remaining HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    
    return text.strip()


## UI Layout ##

# Logo and title
col1, col2 = st.columns([1, 20])
with col1:
    st.image("https://upload.wikimedia.org/wikipedia/en/thumb/e/eb/Heinz_Doofenshmirtz.png/135px-Heinz_Doofenshmirtz.png", width=50)
with col2:
    st.title("📚 Research Roundup-Inator")

# Sidebar for API keys
with st.sidebar:
    st.header("API Configuration")
    use_ollama = st.checkbox("Use Ollama (local)", help="Use local Ollama instance instead of OpenAI (will only work on local machine with Ollama running)")
    if use_ollama:
        st.info("Using Ollama with llama3.1 model")
    else:
        st.session_state.openai_api_key = st.text_input("OpenAI API Key", type="password")
    semanticscholar_api_key = st.text_input("Semantic Scholar API Key (optional)", type="password")
    st.markdown("---")
    st.markdown("### About")
    st.markdown("This tool uses Semantic Scholar's API and LLMs to generate research roundups from academic papers. Launch Ollama or provide your OpenAI API key above to get started. \n\n*(don't worry, your key stays secure on your device)* \n\nMade with ❤️ by [Ben Rochford](https://benrochford.com)")

# Main content
tab1, tab2, tab3 = st.tabs(["Search Papers", "Browse Collected Papers", "Generate Roundup"])

with tab1:
    st.header("Search papers with Semantic Scholar")
    
    # Date controls section
    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        last_2_weeks = st.checkbox("Last 2 weeks", value=False)
    
    # Update dates based on last_2_weeks checkbox
    if last_2_weeks:
        start_date = datetime.now() - timedelta(days=14)
        end_date = datetime.now()
    else:
        with col2:
            start_date = st.date_input("Start Date (optional)", value=None)
        with col3:
            end_date = st.date_input("End Date (optional)", value=None)
    
    search_method = st.radio("Search Method", ["Provide search terms", "Generate search terms automatically"])
    
    # Initialize queries and selected_queries in session state if not exists
    if 'generated_queries' not in st.session_state:
        st.session_state.generated_queries = []
    if 'selected_queries' not in st.session_state:
        st.session_state.selected_queries = {}
    
    # Initialize queries
    queries = []
    
    if search_method == "Provide search terms":
        search_terms = st.text_area(
            "Enter search terms (one per line)",
            help="""Supports advanced search operators:
            - Use quotes for exact phrases: "machine learning"
            - Use + to require terms: +ai +ethics
            - Use - to exclude terms: -privacy
            - Use | for OR: (ai | artificial intelligence)
            - Use * for prefix matching: neural*
            - Use ~N for fuzzy matching: algorithm~2
            - Use "word1 word2" ~N for proximity search
            """
        )
        queries = [term.strip() for term in search_terms.split('\n') if term.strip()]
    else:
        topic_description = st.text_area(
            "Describe your research topic",
            help="Describe the topic you want to research, and the LLM will generate a batch of terms",
            placeholder="Example: Recent advances in quantum computing focusing on error correction"
        )
        
        if topic_description:
            if st.button("Generate Search Terms"):
                if not (st.session_state.openai_api_key or use_ollama):
                    st.warning("⚠️ Please configure LLM in the sidebar first")
                else:
                    with st.spinner("Generating search terms..."):
                        generated_terms = suggest_search_terms(
                            topic_description, 
                            st.session_state.openai_api_key,
                            use_ollama
                        )
                        st.session_state.generated_queries = [
                            term.strip().lstrip("0123456789.- ")
                            for term in generated_terms
                            if term.strip()
                        ]
                        # Initialize all new terms as selected
                        for query in st.session_state.generated_queries:
                            if query not in st.session_state.selected_queries:
                                st.session_state.selected_queries[query] = True
                
                st.write("Generated search terms:")
                for i, query in enumerate(st.session_state.generated_queries):
                    col1, col2 = st.columns([1, 20])
                    with col1:
                        st.session_state.selected_queries[query] = st.checkbox(
                            "Use",
                            value=st.session_state.selected_queries.get(query, True),
                            key=f"search_term_checkbox_{i}",
                            label_visibility="collapsed"
                        )
                    with col2:
                        st.code(query)
                
                # Use only the selected generated queries for the search
                queries = [q for q in st.session_state.generated_queries 
                          if st.session_state.selected_queries.get(q, True)]
            else:
                # Show previously generated queries with their checkboxes if they exist
                if st.session_state.generated_queries:
                    st.write("Generated search terms:")
                    for i, query in enumerate(st.session_state.generated_queries):
                        col1, col2 = st.columns([1, 20])
                        with col1:
                            st.session_state.selected_queries[query] = st.checkbox(
                                "Use",
                                value=st.session_state.selected_queries.get(query, True),
                                key=f"search_term_checkbox_{i}",
                                label_visibility="collapsed"
                            )
                        with col2:
                            st.code(query)
                
                    # Use only the selected generated queries for the search
                    queries = [q for q in st.session_state.generated_queries 
                              if st.session_state.selected_queries.get(q, True)]
    
    # Update excluded terms input and processing
    excluded_terms_input = st.text_area(
        "Excluded terms (one per line)",
        help="""Each term will be automatically prefixed with - to exclude it.
        For exact phrases, use quotes: "machine learning"
        """,
        placeholder="Enter terms to exclude, one per line"
    )
    # Process excluded terms as complete lines, not individual characters
    excluded_terms = [line.strip() for line in excluded_terms_input.split('\n') if line.strip()]

    # Preview API call
    if queries:
        st.write("---")
        st.write("Preview of API calls to be made:")
        base_url = "https://api.semanticscholar.org/graph/v1/paper/search"
        params = {
            "fields": "title,url,abstract,citationCount,authors,publicationTypes,publicationDate,openAccessPdf",
            "limit": 100
        }
        
        # Add date parameters if applicable
        if last_2_weeks:
            end_date = pd.Timestamp.now()
            start_date = end_date - pd.Timedelta(days=14)
            params["publicationDateOrYear"] = f"{start_date.strftime('%Y-%m-%d')}:{end_date.strftime('%Y-%m-%d')}"
        elif start_date and end_date:  # Only add year parameter if both dates are provided
            params["year"] = f"{start_date.year}-{end_date.year}"
        elif start_date:  # Only start date provided
            params["year"] = f"{start_date.year}-"
        elif end_date:  # Only end date provided
            params["year"] = f"-{end_date.year}"
        
        # Show URL for first query as example
        example_query = queries[0]
        if excluded_terms:
            excluded_query = ' '.join(f'-{term}' for term in excluded_terms)
            example_query = f'{example_query} {excluded_query}'
        
        params["query"] = example_query
        preview_url = f"{base_url}?{'&'.join(f'{k}={requests.utils.quote(str(v))}' for k, v in params.items())}"
        st.code(preview_url, language="text")
        if len(queries) > 1:
            st.caption(f"(+ {len(queries)-1} more queries)")
    
    # Update search button section
    if queries:
        st.write("---")
        col1, col2 = st.columns([1, 3])
        with col1:
            search_clicked = st.button(
                "🚀 Search Semantic Scholar",
                disabled=not (st.session_state.openai_api_key or use_ollama),
                help="Click to search for papers on Semantic Scholar using the selected terms"
            )
        with col2:
            if not (st.session_state.openai_api_key or use_ollama):
                st.warning("⚠️ Please enter your OpenAI API key in the sidebar first")
        
        if search_clicked:
            with st.spinner("Initializing search..."):
                papers, results_by_term = collect_papers(
                    queries,
                    start_date.year if start_date else None,
                    end_date.year if end_date else None,
                    last_two_weeks=last_2_weeks,
                    api_key=semanticscholar_api_key,
                    excluded_terms=excluded_terms
                )
                st.session_state.papers_df = process_papers(papers)
                
                # Success message with tab navigation instruction
                st.success(f"✨ Search completed! Found {len(st.session_state.papers_df)} unique papers")
                st.info("👆 Click the 'Browse Collected Papers' tab above to review the results")

with tab2:
    if st.session_state.papers_df is not None:
        st.header("Browse Collected Papers", anchor="review-collected-papers")
        st.dataframe(
            st.session_state.papers_df,
            column_config={
                "title": st.column_config.TextColumn("Title", width="large"),
                "url": st.column_config.LinkColumn("URL"),
                "abstract": st.column_config.TextColumn("Abstract", width="large"),
            },
            hide_index=True,
        )
    else:
        st.info("No papers collected yet. Use the Search tab to find papers.")

with tab3:
    if st.session_state.papers_df is not None:
        st.header("Generate Roundup")

        # Use columns to make the number input narrower
        col1, col2 = st.columns([1, 3])
        with col1:
            st.session_state.top_n_papers = st.number_input(
                "Number of papers to include in roundup",
                min_value=5,
                max_value=100,
                value=st.session_state.top_n_papers,
                help="Select how many of the most relevant papers to include in the roundup"
            )

        # Compute and display top N papers with relevancy scores
        st.session_state.papers_df['score'] = st.session_state.papers_df['citationCount'].fillna(0) + \
            100*(1 / (1 + (pd.Timestamp.now() - st.session_state.papers_df['publicationDate']).dt.days))
        
        top_papers_df = st.session_state.papers_df.sort_values('score', ascending=False).head(st.session_state.top_n_papers)
        top_papers_df = top_papers_df[['score'] + [col for col in top_papers_df.columns if col != 'score']]
        
        with st.expander("📄 Papers Selected for Roundup", expanded=False):
            st.write(f"The top {st.session_state.top_n_papers} most relevant papers selected for the roundup, based on recency and citation count")
            st.dataframe(top_papers_df)
        
        # Add prompt customization with expander
        with st.expander("📝 Customize Generation Prompt", expanded=False):
            
            col1, col2 = st.columns([3, 10])
            with col1:
                st.markdown("#### Roundup Prompt")
            with col2:
                if st.button("reset to default", type="secondary", key="reset_prompt"):
                    st.session_state.custom_prompt = DEFAULT_ROUNDUP_PROMPT
            
            custom_prompt = st.text_area(
                "Generation Prompt",
                value=st.session_state.custom_prompt,
                height=400,
                help="Edit this prompt to customize how your research roundup is generated",
                label_visibility="collapsed"
            )
            
            # Save any changes to the prompt
            if custom_prompt != st.session_state.custom_prompt:
                st.session_state.custom_prompt = custom_prompt
        
        if st.button("Generate Research Roundup") and (st.session_state.openai_api_key or use_ollama):
            progress_placeholder = st.empty()
            message_placeholder = st.empty()
            
            with progress_placeholder:
                with st.spinner("Generating roundup..."):
                    # Start newsletter generation in background
                    newsletter_content = None
                    start_time = time.time()
                    
                    # Keep showing messages until generation is complete
                    message_index = 0
                    while newsletter_content is None:
                        # Cycle through messages
                        dots = "." * ((message_index % 3) + 1)
                        message = LOADING_MESSAGES[message_index % len(LOADING_MESSAGES)]
                        message_placeholder.markdown(f"*{message}{dots}*")
                        
                        # Generate newsletter on first iteration
                        if message_index == 0:
                            newsletter_content = generate_newsletter(
                                st.session_state.papers_df,
                                st.session_state.openai_api_key,
                                st.session_state.custom_prompt,
                                use_ollama,
                                st.session_state.top_n_papers
                            )
                        
                        message_index += 1
                        time.sleep(0.8)
                    
                    st.session_state.newsletter_content = newsletter_content
            
            # Clear the loading message
            message_placeholder.empty()
        if st.session_state.newsletter_content:
            st.markdown("---")
            st.markdown(st.session_state.newsletter_content)
            st.markdown("---")

            # Download buttons
            col1, col2, col3, col4 = st.columns([4,3,3,4])
            with col2:
                st.download_button(
                    "Download as Markdown", 
                    st.session_state.newsletter_content,
                    file_name="research_roundup.md",
                    mime="text/markdown"
                )
            with col3:
                docx_file = convert_to_docx(st.session_state.newsletter_content)
                st.download_button(
                    "Download as Word",
                    docx_file, 
                    file_name="research_roundup.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.info("No papers collected yet. Use the Search tab to find papers.")
